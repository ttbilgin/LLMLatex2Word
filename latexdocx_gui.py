import sys
import os
import tempfile
import subprocess
from tkinter import Tk, Text, Button, filedialog
from docx import Document
from lxml import etree
import latex2mathml.converter
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt  # For setting font size


def clean_latex_formula(formula):
    """LaTeX formülünü temizler ve dönüştürmeye hazırlar"""
    # Delimiterleri kaldır
    if formula.startswith('$') and formula.endswith('$'):
        formula = formula[1:-1]
    elif formula.startswith('$$') and formula.endswith('$$'):
        formula = formula[2:-2]
    elif formula.startswith('\\[') and formula.endswith('\\]'):
        formula = formula[2:-2]
    elif formula.startswith('\\(') and formula.endswith('\\)'):
        formula = formula[2:-2]
    
    # Gereksiz boşlukları temizle
    formula = formula.strip()
    return formula

def latex_to_word(latex_input):
    try:
        # LaTeX formülünü temizle
        cleaned_latex = clean_latex_formula(latex_input)
        
        # MathML'e dönüştür
        mathml = latex2mathml.converter.convert(cleaned_latex)
        
        # XML ağacını oluştur
        tree = etree.fromstring(mathml)
        
        # XSLT dönüşümünü uygula
        xslt = etree.parse('MML2OMML.XSL')
        transform = etree.XSLT(xslt)
        new_dom = transform(tree)
        
        return new_dom.getroot()
    except Exception as e:
        print(f"Error converting formula: {latex_input}")
        print(f"Error message: {str(e)}")
        return None

def process_text_with_bold(paragraph, text):
    # \textbf{...} pattern'ını bul
    bold_pattern = re.compile(r'\\textbf\{([^}]+)\}')
    
    last_end = 0
    for match in bold_pattern.finditer(text):
        # Normal metni ekle
        normal_text = text[last_end:match.start()].strip()
        if normal_text:
            paragraph.add_run(normal_text)
            paragraph.add_run(" ")  # Bold metinden önce boşluk
        
        # Kalın metni ekle
        bold_text = match.group(1).strip()  # Süslü parantezler içindeki metin
        run = paragraph.add_run(bold_text)
        run.bold = True
        paragraph.add_run(" ")  # Bold metinden sonra boşluk
        
        last_end = match.end()
    
    # Kalan normal metni ekle
    remaining_text = text[last_end:].strip()
    if remaining_text:
        paragraph.add_run(remaining_text)

def add_latex_and_text_to_docx(document, content):
    latex_pattern = re.compile(r'''
        (?:
            \$\$.*?\$\$               # Display math with $$..$$
            |\$[^$]+?\$               # Inline math with $..$
            |\\\[[\s\S]*?\\\]        # Display math with \[..\]
            |\\\([\s\S]*?\\\)        # Inline math with \(..\)
        )
    ''', re.VERBOSE | re.DOTALL)
    
    parts = []
    last_end = 0
    
    for match in latex_pattern.finditer(content):
        text_before = content[last_end:match.start()].strip()
        formula = match.group()
        
        if text_before:
            parts.append(("text", text_before))
        parts.append(("formula", formula))
        last_end = match.end()
    
    text_after = content[last_end:].strip()
    if text_after:
        parts.append(("text", text_after))
    
    current_paragraph = None
    
    for i, (part_type, part_content) in enumerate(parts):
        if part_type == "formula":
            cleaned_formula = clean_latex_formula(part_content)
            is_short_formula = len(cleaned_formula) <= 50
            
            if is_short_formula:
                if current_paragraph is None:
                    current_paragraph = document.add_paragraph()
                
                current_paragraph.add_run(" ")
                
                word_math = latex_to_word(part_content)
                if word_math is not None:
                    current_paragraph._element.append(word_math)
                
                if i < len(parts) - 1 and parts[i + 1][0] == "text":
                    current_paragraph.add_run(" ")
                    continue
            else:
                current_paragraph = document.add_paragraph()
                word_math = latex_to_word(part_content)
                if word_math is not None:
                    current_paragraph._element.append(word_math)
                current_paragraph = None
        else:  # text
            if current_paragraph is None:
                current_paragraph = document.add_paragraph()
            # Normal metin ekleme yerine process_text_with_bold kullan
            process_text_with_bold(current_paragraph, part_content)
            
            if i == len(parts) - 1 or (parts[i + 1][0] == "formula" and 
                len(clean_latex_formula(parts[i + 1][1])) > 3):
                current_paragraph = None

def create_docx_from_tex(tex_content):
    # Create a Word document
    document = Document()

    # Add content
    add_latex_and_text_to_docx(document, tex_content)
	
    # Set justification and font style for all paragraphs
    for paragraph in document.paragraphs:
        # Set alignment to justify
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in paragraph.runs:
            run.font.name = "Times New Roman"  # Set font to Times New Roman
            run.font.size = Pt(11)  # Set font size to 12pt (or any size you prefer)

    # Save to a temporary file
    temp_dir = tempfile.gettempdir()
    output_path = os.path.join(temp_dir, 'output.docx')
    document.save(output_path)
    return output_path

def open_word_document(docx_path):
    # Open the document in MS Word
    try:
        subprocess.run(['start', docx_path], shell=True, check=True)
    except Exception as e:
        print(f"Error opening Word document: {e}")

def on_run_button_click():
    tex_content = text_widget.get("1.0", "end-1c")
    if not tex_content.strip():
        print("No content provided.")
        return

    docx_path = create_docx_from_tex(tex_content)
    print(f"Document saved at {docx_path}")
    open_word_document(docx_path)

# Create the GUI
root = Tk()
root.title("LaTeX to Word Converter")

# Text widget for entering LaTeX content
text_widget = Text(root, wrap="word", width=150, height=40)
text_widget.pack(padx=10, pady=10)

# Button to generate the Word document
run_button = Button(root, text="CONVERT TO MSWORD", command=on_run_button_click)
run_button.pack(pady=10)

root.mainloop()
